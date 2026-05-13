import os
from docx import Document
from docx.shared import Pt


def merge_py_files_to_word(source_folder, output_docx_path):
    # Создаем новый документ Word
    doc = Document()

    # Проверяем, существует ли указанная папка
    if not os.path.exists(source_folder):
        print(f"Ошибка: Папка '{source_folder}' не найдена.")
        return

    # Получаем список всех файлов в папке и фильтруем только .py
    files = [f for f in os.listdir(source_folder) if f.endswith(".py")]

    if not files:
        print(f"В папке '{source_folder}' не найдено файлов с расширением .py.")
        return

    print(f"Найдено файлов для переноса: {len(files)}")

    # Обрабатываем каждый файл
    for file_name in sorted(files):
        file_path = os.path.join(source_folder, file_name)

        # Добавляем название файла как заголовок в документ
        doc.add_heading(f"Файл: {file_name}", level=2)

        try:
            # Читаем содержимое .py файла
            with open(file_path, "r", encoding="utf-8") as f:
                code_content = f.read()

            # Добавляем код в документ
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(code_content)

            # Настраиваем моноширинный шрифт для сохранения отступов (табуляции)
            run.font.name = "Courier New"
            run.font.size = Pt(10.5)

            # Добавляем разрыв страницы после каждого файла (опционально)
            doc.add_page_break()
            print(f"Успешно добавлен: {file_name}")

        except Exception as e:
            print(f"Не удалось прочитать файл {file_name}. Ошибка: {e}")

    # Сохраняем итоговый документ
    doc.save(output_docx_path)
    print(f"\nГотово! Все файлы сохранены в: {output_docx_path}")


# --- НАСТРОЙКА ПУТЕЙ ---
# Укажите путь к папке, где лежат ваши .py файлы
TARGET_DIR = "./modules"
# Укажите имя итогового файла Word
RESULT_FILE = "all_scripts_documentation.docx"

# Запуск функции
merge_py_files_to_word(TARGET_DIR, RESULT_FILE)
